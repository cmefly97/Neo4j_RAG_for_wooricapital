"""
형식별 파서 — 문서를 '청크' 리스트로 변환.
각 청크: {text, source_file, locator, chunk_id}

- md   : 헤더(#) 기준 섹션 청킹
- docx : 문단/표 (Heading 인지)
- pdf  : 페이지 + 섹션 분할
- xlsx : 시트 → 행 레코드(텍스트화)
"""
from __future__ import annotations
import hashlib
import re
from pathlib import Path


def _cid(source_file: str, locator: str, text: str) -> str:
    h = hashlib.md5(f"{source_file}|{locator}|{text[:80]}".encode("utf-8")).hexdigest()[:10]
    return f"{Path(source_file).stem[:20]}::{locator}::{h}"


def _mk(text, source_file, locator):
    text = (text or "").strip()
    return {"text": text, "source_file": source_file, "locator": locator,
            "chunk_id": _cid(source_file, locator, text)}


def parse(path: str) -> list[dict]:
    ext = Path(path).suffix.lower()
    if ext == ".md":   return parse_md(path)
    if ext == ".docx": return parse_docx(path)
    if ext == ".pdf":  return parse_pdf(path)
    if ext == ".xlsx": return parse_xlsx(path)
    raise ValueError(f"지원하지 않는 형식: {ext}")


def parse_md(path: str) -> list[dict]:
    name = Path(path).name
    lines = Path(path).read_text(encoding="utf-8").splitlines()
    chunks, buf, heading = [], [], "intro"
    def flush():
        if buf:
            body = "\n".join(buf).strip()
            if body:
                chunks.append(_mk(f"# {heading}\n{body}", name, heading[:60]))
    for ln in lines:
        m = re.match(r"^(#{1,4})\s+(.*)", ln)
        if m:
            flush(); buf = []
            heading = m.group(2).strip()
        else:
            buf.append(ln)
    flush()
    return [c for c in chunks if len(c["text"]) > 10]


def parse_docx(path: str) -> list[dict]:
    import docx
    name = Path(path).name
    doc = docx.Document(path)
    chunks, buf, heading = [], [], "intro"
    def flush():
        if buf:
            body = "\n".join(buf).strip()
            if body:
                chunks.append(_mk(f"# {heading}\n{body}", name, heading[:60]))
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        if (p.style and p.style.name and p.style.name.lower().startswith("heading")):
            flush(); buf = []
            heading = t
        else:
            buf.append(t)
    flush()
    # 표는 행 단위로 텍스트화해 별도 청크
    for ti, tbl in enumerate(doc.tables):
        rows = [" | ".join(c.text.strip() for c in r.cells) for r in tbl.rows]
        body = "\n".join(r for r in rows if r.strip(" |"))
        if body:
            chunks.append(_mk(f"[표{ti+1}]\n{body}", name, f"table{ti+1}"))
    return [c for c in chunks if len(c["text"]) > 10]


def parse_pdf(path: str) -> list[dict]:
    import fitz  # pymupdf
    name = Path(path).name
    chunks = []
    with fitz.open(path) as doc:
        for pno in range(len(doc)):
            text = doc[pno].get_text("text").strip()
            if not text:
                continue
            # 긴 페이지는 빈 줄 2개 이상 기준으로 분할
            parts = re.split(r"\n\s*\n\s*\n+", text)
            for si, part in enumerate(parts):
                part = part.strip()
                if len(part) > 15:
                    loc = f"p{pno+1}" + (f".{si+1}" if len(parts) > 1 else "")
                    chunks.append(_mk(part, name, loc))
    return chunks


def parse_xlsx(path: str) -> list[dict]:
    import pandas as pd
    name = Path(path).name
    chunks = []
    xl = pd.ExcelFile(path)
    for sheet in xl.sheet_names:
        df = xl.parse(sheet, dtype=str).fillna("")
        cols = list(df.columns)
        for i, row in df.iterrows():
            cells = [f"{c}: {str(row[c]).strip()}" for c in cols if str(row[c]).strip()]
            body = " | ".join(cells)
            if len(body) > 5:
                chunks.append(_mk(body, name, f"{sheet}#row{i+1}"))
    return chunks
